#!/usr/bin/env python3
"""Convert constrained Qwen manuscript Markdown to DOCX with true footnotes."""

from __future__ import annotations

import argparse
import re
import zipfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
PKGREL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
NS = {"w": W_NS}
BODY_FONT = "新細明體"
HEADING_FONT = "微軟正黑體"


@dataclass
class Block:
    kind: str
    text: str


def parse_markdown(text: str) -> tuple[list[Block], dict[int, tuple[str, bool]]]:
    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    notes: dict[int, tuple[str, bool]] = {}
    body_lines: list[str] = []
    current_note: int | None = None

    for line in lines:
        match = re.match(r"^\[\^(\d+)\]:\s*(.*)$", line)
        if match:
            number = int(match.group(1))
            if number in notes:
                raise ValueError(f"duplicate footnote definition: {number}")
            content = match.group(2).strip()
            italic = content.startswith("*") and content.endswith("*") and len(content) >= 2
            if italic:
                content = content[1:-1].strip()
            notes[number] = (content, italic or ("原稿" in content and "待補" in content))
            current_note = number
        elif current_note is not None and (line.startswith("    ") or line.startswith("\t")):
            old, italic = notes[current_note]
            notes[current_note] = (old + " " + line.strip(), italic)
        else:
            current_note = None
            body_lines.append(line)

    blocks: list[Block] = []
    buffer: list[str] = []
    quote_buffer: list[str] = []

    def flush_body() -> None:
        if buffer:
            blocks.append(Block("paragraph", "".join(part.strip() for part in buffer)))
            buffer.clear()

    def flush_quote() -> None:
        if quote_buffer:
            blocks.append(Block("quote", "".join(part.strip() for part in quote_buffer)))
            quote_buffer.clear()

    for line in body_lines:
        stripped = line.strip()
        if not stripped:
            flush_body()
            flush_quote()
        elif stripped.startswith("### "):
            flush_body()
            flush_quote()
            blocks.append(Block("heading", stripped[4:].strip()))
        elif stripped.startswith(">"):
            flush_body()
            quote_buffer.append(stripped[1:].lstrip())
        else:
            flush_quote()
            buffer.append(line)
    flush_body()
    flush_quote()

    if not blocks or blocks[0].kind != "heading":
        raise ValueError("Markdown must begin with a level-3 heading (###)")

    refs = [int(n) for block in blocks for n in re.findall(r"\[\^(\d+)\]", block.text)]
    if len(refs) != len(set(refs)):
        raise ValueError("each footnote label must be referenced exactly once")
    if set(refs) != set(notes):
        missing_defs = sorted(set(refs) - set(notes))
        unused_defs = sorted(set(notes) - set(refs))
        raise ValueError(f"footnote mismatch: missing definitions={missing_defs}, unused definitions={unused_defs}")
    if refs:
        expected = list(range(refs[0], refs[0] + len(refs)))
        if refs != expected:
            raise ValueError(f"footnotes must be contiguous and ordered by first use: got {refs}, expected {expected}")
    return blocks, notes


def set_run_font(run, name: str, size: float, bold: bool = False, italic: bool = False) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)


def set_style_font(style, name: str, size: float, bold: bool = False) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)


def configure_paragraph(paragraph, quote: bool = False) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Pt(0 if quote else 24)
    fmt.left_indent = Pt(24 if quote else 0)
    fmt.right_indent = Pt(24 if quote else 0)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(8)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = 1.333


INLINE_RE = re.compile(r"(\[\^\d+\]|\*\*.+?\*\*|(?<!\*)\*[^*]+?\*(?!\*))")


def add_inline(paragraph, text: str) -> None:
    position = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            set_run_font(run, BODY_FONT, 11)
        token = match.group(0)
        if token.startswith("[^" ):
            run = paragraph.add_run(f"[[FN{token[2:-1]}]]")
            set_run_font(run, BODY_FONT, 11)
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, BODY_FONT, 11, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, BODY_FONT, 11, italic=True)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, BODY_FONT, 11)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    set_run_font(run, "Calibri", 9)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, end))


def build_base(blocks: list[Block], path: Path, header_text: str) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    set_style_font(normal, BODY_FONT, 11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    heading = doc.styles["Heading 1"]
    set_style_font(heading, HEADING_FONT, 16, True)
    heading.paragraph_format.space_before = Pt(18)
    heading.paragraph_format.space_after = Pt(10)
    heading.paragraph_format.keep_with_next = True
    try:
        footnote_style = doc.styles["Footnote Text"]
    except KeyError:
        footnote_style = doc.styles.add_style("Footnote Text", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(footnote_style, BODY_FONT, 9)
    footnote_style.paragraph_format.space_after = Pt(2)

    header = section.header.paragraphs[0]
    run = header.add_run(header_text)
    set_run_font(run, HEADING_FONT, 9)
    run.font.color.rgb = RGBColor(128, 128, 128)
    add_page_field(section.footer.paragraphs[0])

    for block in blocks:
        if block.kind == "heading":
            paragraph = doc.add_paragraph(style="Heading 1")
            paragraph.add_run(block.text)
        else:
            paragraph = doc.add_paragraph()
            configure_paragraph(paragraph, quote=block.kind == "quote")
            add_inline(paragraph, block.text)
    doc.save(path)


def xml_bytes(root: etree._Element) -> bytes:
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")


def make_footnotes_part() -> etree._Element:
    root = etree.Element(f"{{{W_NS}}}footnotes", nsmap={"w": W_NS, "r": R_NS})
    for note_id, tag in ((-1, "separator"), (0, "continuationSeparator")):
        note = etree.SubElement(root, f"{{{W_NS}}}footnote")
        note.set(f"{{{W_NS}}}id", str(note_id))
        paragraph = etree.SubElement(note, f"{{{W_NS}}}p")
        run = etree.SubElement(paragraph, f"{{{W_NS}}}r")
        etree.SubElement(run, f"{{{W_NS}}}{tag}")
    return root


def set_superscript(run: etree._Element) -> None:
    rpr = run.find("w:rPr", namespaces=NS)
    if rpr is None:
        rpr = etree.Element(f"{{{W_NS}}}rPr")
        run.insert(0, rpr)
    vert = etree.SubElement(rpr, f"{{{W_NS}}}vertAlign")
    vert.set(f"{{{W_NS}}}val", "superscript")


def append_note(root: etree._Element, note_id: int, text: str, italic: bool) -> None:
    note = etree.SubElement(root, f"{{{W_NS}}}footnote")
    note.set(f"{{{W_NS}}}id", str(note_id))
    paragraph = etree.SubElement(note, f"{{{W_NS}}}p")
    ppr = etree.SubElement(paragraph, f"{{{W_NS}}}pPr")
    pstyle = etree.SubElement(ppr, f"{{{W_NS}}}pStyle")
    pstyle.set(f"{{{W_NS}}}val", "FootnoteText")
    ref_run = etree.SubElement(paragraph, f"{{{W_NS}}}r")
    set_superscript(ref_run)
    etree.SubElement(ref_run, f"{{{W_NS}}}footnoteRef")
    text_run = etree.SubElement(paragraph, f"{{{W_NS}}}r")
    rpr = etree.SubElement(text_run, f"{{{W_NS}}}rPr")
    fonts = etree.SubElement(rpr, f"{{{W_NS}}}rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(f"{{{W_NS}}}{attr}", BODY_FONT)
    size = etree.SubElement(rpr, f"{{{W_NS}}}sz")
    size.set(f"{{{W_NS}}}val", "18")
    if italic:
        etree.SubElement(rpr, f"{{{W_NS}}}i")
    node = etree.SubElement(text_run, f"{{{W_NS}}}t")
    node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    node.text = " " + text


def patch_docx(base: Path, output: Path, notes: dict[int, tuple[str, bool]]) -> None:
    ordered = sorted(notes)
    internal_ids = {label: index for index, label in enumerate(ordered, start=1)}
    with zipfile.ZipFile(base, "r") as zin:
        doc_root = etree.fromstring(zin.read("word/document.xml"))
        footnotes_root = make_footnotes_part()

        for label in ordered:
            marker = f"[[FN{label}]]"
            found = False
            for text_node in doc_root.xpath(".//w:t", namespaces=NS):
                if text_node.text == marker:
                    run = text_node.getparent()
                    parent = run.getparent()
                    index = parent.index(run)
                    parent.remove(run)
                    ref_run = etree.Element(f"{{{W_NS}}}r")
                    set_superscript(ref_run)
                    ref = etree.SubElement(ref_run, f"{{{W_NS}}}footnoteReference")
                    ref.set(f"{{{W_NS}}}id", str(internal_ids[label]))
                    parent.insert(index, ref_run)
                    found = True
                    break
            if not found:
                raise RuntimeError(f"internal footnote marker missing: {label}")
            content, italic = notes[label]
            append_note(footnotes_root, internal_ids[label], content, italic)

        sect_pr = doc_root.find(".//w:sectPr", namespaces=NS)
        footnote_pr = etree.Element(f"{{{W_NS}}}footnotePr")
        num_start = etree.SubElement(footnote_pr, f"{{{W_NS}}}numStart")
        num_start.set(f"{{{W_NS}}}val", str(ordered[0] if ordered else 1))
        num_restart = etree.SubElement(footnote_pr, f"{{{W_NS}}}numRestart")
        num_restart.set(f"{{{W_NS}}}val", "continuous")
        sect_pr.insert(0, footnote_pr)

        rels = etree.fromstring(zin.read("word/_rels/document.xml.rels"))
        rel_ids = [int(m.group(1)) for rel in rels for m in [re.match(r"rId(\d+)$", rel.get("Id", ""))] if m]
        rel = etree.SubElement(rels, f"{{{PKGREL_NS}}}Relationship")
        rel.set("Id", f"rId{max(rel_ids, default=0) + 1}")
        rel.set("Type", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes")
        rel.set("Target", "footnotes.xml")

        content_types = etree.fromstring(zin.read("[Content_Types].xml"))
        override = etree.SubElement(content_types, f"{{{CT_NS}}}Override")
        override.set("PartName", "/word/footnotes.xml")
        override.set("ContentType", "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml")

        replacements = {
            "word/document.xml": xml_bytes(doc_root),
            "word/footnotes.xml": xml_bytes(footnotes_root),
            "word/_rels/document.xml.rels": xml_bytes(rels),
            "[Content_Types].xml": xml_bytes(content_types),
        }
        with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as zout:
            existing = set()
            for item in zin.infolist():
                existing.add(item.filename)
                zout.writestr(item, replacements.get(item.filename, zin.read(item.filename)))
            for name, data in replacements.items():
                if name not in existing:
                    zout.writestr(name, data)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--header", default="學術書稿｜千問潤稿")
    args = parser.parse_args()

    blocks, notes = parse_markdown(args.input.read_text(encoding="utf-8-sig"))
    args.output.parent.mkdir(parents=True, exist_ok=True)
    base = args.output.with_suffix(".base.docx")
    build_base(blocks, base, args.header)
    patch_docx(base, args.output, notes)
    base.unlink(missing_ok=True)
    start = min(notes) if notes else None
    end = max(notes) if notes else None
    print(f"created={args.output}")
    print(f"footnotes={len(notes)} range={start}-{end}")


if __name__ == "__main__":
    main()
